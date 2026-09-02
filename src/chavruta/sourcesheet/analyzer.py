"""Source Sheet Companion Analyzer & Synthesizer (Spec 008 Phase 3).

Grounded sugya flow analysis, XML-isolated source prompt construction,
Mermaid flowchart generation, comparative opinion tables, tiered chavruta
questions, and multi-format export (Markdown, Word .docx).
"""

from __future__ import annotations

import html
import io
import json
import logging
import re
from dataclasses import asdict, dataclass, field
from typing import Any

from chavruta.corpus.schema import Citation
from chavruta.llm.base import GroundedPrompt
from chavruta.sourcesheet.parser import ParsedSourceItem

_log = logging.getLogger("chavruta.sourcesheet.analyzer")

# ── Verification Statuses (Principle I) ──────────────────────────────────────

STATUS_VERIFIED_CORPUS = "VERIFIED_CORPUS"
STATUS_USER_PROVIDED = "USER_PROVIDED"
STATUS_MISSING_REF = "MISSING_REF"

STATUS_LABELS_HE = {
    STATUS_VERIFIED_CORPUS: "מאומת ומורחב מהמאגר",
    STATUS_USER_PROVIDED: "מקור מדף המשתמש — לא אומת מול מאגר חברותא",
    STATUS_MISSING_REF: "מראה מקום שלא נמצא במאגר (חסר טקסט)",
}


@dataclass
class SourceSection:
    index: int
    title: str
    ref: str | None
    status: str
    role_tag: str
    source_snippet: str
    expanded_context: str | None = None
    plain_explanation: str = ""
    diyuk: str | None = None
    difficult_words: dict[str, str] = field(default_factory=dict)
    author_note: str | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class CompanionGuide:
    title: str
    topic: str
    core_inquiry: str
    flowchart_mermaid: str
    sections: list[SourceSection]
    opinion_table: list[dict[str, str]]
    chavruta_questions: dict[str, list[str]]  # peshat, comparison, sevara
    summary: str
    citations: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "title": self.title,
            "topic": self.topic,
            "core_inquiry": self.core_inquiry,
            "flowchart_mermaid": self.flowchart_mermaid,
            "sections": [s.to_dict() for s in self.sections],
            "opinion_table": self.opinion_table,
            "chavruta_questions": self.chavruta_questions,
            "summary": self.summary,
            "citations": self.citations,
        }

    def to_markdown(self) -> str:
        """Produce a complete, beautifully-formatted Markdown companion guide."""
        from chavruta.corpus.refs import hebrew_display_ref

        md = []
        md.append(f"# חוברת ליווי לדף מקורות: {self.title or self.topic}\n")
        md.append(f"> **נושא הסוגיה:** {self.topic}\n")
        md.append("## שאלת היסוד וציר החקירה\n")
        md.append(f"{self.core_inquiry}\n")

        if self.flowchart_mermaid:
            md.append("## מפת מהלך הסוגיה\n")
            md.append("```mermaid")
            md.append(self.flowchart_mermaid)
            md.append("```\n")

        md.append("## ביאור מקורות הדף\n")
        for s in self.sections:
            display_ref = (hebrew_display_ref(s.ref) or s.ref) if s.ref else ""
            md.append(f"### מקור {s.index}: {s.title} ({s.role_tag})\n")
            if display_ref:
                status_badge = "מאומת מהמאגר" if s.status == STATUS_VERIFIED_CORPUS else "מדף המקורות"
                md.append(f"**מראה מקום:** {display_ref} *({status_badge})*\n")

            if s.status == STATUS_MISSING_REF:
                md.append(
                    "> ⚠️ **הערת מערכת:** מראה מקום זה אינו מצוטט בדף ואינו קיים במאגר. "
                    "לא ניתן לבאר ללא ציטוט הטקסט — באפשרותך להזין את לשונו בצ'אט.\n"
                )
                continue

            md.append(f"**לשון המקור:**\n> {s.source_snippet}\n")
            if s.plain_explanation:
                md.append(f"**ביאור הפשט והמהלך:**\n{s.plain_explanation}\n")
            if s.diyuk:
                md.append(f"**נקודת הדיוק:**\n*{s.diyuk}*\n")
            if s.difficult_words:
                words_str = ", ".join(f"**{k}**: {v}" for k, v in s.difficult_words.items())
                md.append(f"**ביאורי מילים:** {words_str}\n")
            if s.author_note:
                md.append(f"💡 **הערת עורך הדף:** {s.author_note}\n")
            md.append("---\n")

        if self.opinion_table:
            md.append("## טבלת השוואת שיטות ומחלוקות\n")
            md.append("| שיטה / בעל הדעה | סברת היסוד | הראיה / המקור | נפקא מינה |")
            md.append("|---|---|---|---|")
            for row in self.opinion_table:
                opinion = row.get("opinion", "")
                reason = row.get("reason", "")
                proof = row.get("proof", "")
                nafka = row.get("nafka_mina", "")
                md.append(f"| {opinion} | {reason} | {proof} | {nafka} |")
            md.append("")

        if self.chavruta_questions:
            md.append("## שאלות לעיון וחזרת החברותא\n")
            if self.chavruta_questions.get("peshat"):
                md.append("### 1. שאלות פשט והבנת המקורות:")
                for q in self.chavruta_questions["peshat"]:
                    md.append(f"- {q}")
            if self.chavruta_questions.get("comparison"):
                md.append("\n### 2. שאלות השוואה ודיוק שיטות:")
                for q in self.chavruta_questions["comparison"]:
                    md.append(f"- {q}")
            if self.chavruta_questions.get("sevara"):
                md.append("\n### 3. שאלות סברא והעמקה מושגית:")
                for q in self.chavruta_questions["sevara"]:
                    md.append(f"- {q}")
            md.append("")

        if self.summary:
            md.append("## סיכום ומסקנות הסוגיה\n")
            md.append(f"{self.summary}\n")

        return "\n".join(md)

    def to_docx_bytes(self) -> bytes:
        """Produce a real, highly-styled Word .docx file with RTL tables, Torah fonts, and headers."""
        try:
            import docx
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from chavruta.corpus.refs import hebrew_display_ref

            doc = docx.Document()

            # Set margins to 2 cm (approx 0.8 in)
            for section in doc.sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            TEKHELET = RGBColor(0, 32, 69)      # #002045
            NAVY_ACCENT = RGBColor(28, 53, 94)  # #1C355E
            GOLD = RGBColor(184, 134, 11)       # #B8860B
            CHARCOAL = RGBColor(28, 26, 23)     # #1C1A17
            MUTED = RGBColor(95, 105, 120)      # Slate muted
            DARK_RED = RGBColor(160, 30, 30)

            FONT_PRIMARY = "Frank Ruhl Libre"
            FONT_FALLBACK = "David"

            def _apply_rtl_p(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, line_spacing=1.15):
                p.alignment = align
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.line_spacing = line_spacing
                pPr = p._p.get_or_add_pPr()
                if pPr.find(qn("w:bidi")) is None:
                    pPr.append(OxmlElement("w:bidi"))

            def _add_run_rtl(p, text, size_pt=11, bold=False, italic=False, color_rgb=None, font_name=FONT_PRIMARY):
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(size_pt)
                if color_rgb:
                    run.font.color.rgb = color_rgb
                run.font.name = font_name
                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rFonts.set(qn("w:cs"), FONT_FALLBACK)
                rFonts.set(qn("w:eastAsia"), FONT_FALLBACK)
                if rPr.find(qn("w:rtl")) is None:
                    rPr.append(OxmlElement("w:rtl"))
                return run

            def _add_heading_1(title_text):
                p = doc.add_paragraph()
                _apply_rtl_p(p, space_before=16, space_after=6)
                _add_run_rtl(p, title_text, size_pt=15, bold=True, color_rgb=TEKHELET)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "12")  # 1.5 pt
                bottom.set(qn("w:space"), "4")
                bottom.set(qn("w:color"), "B8860B")
                pBdr.append(bottom)
                pPr.append(pBdr)

            def _add_heading_2(subtitle_text, color=NAVY_ACCENT):
                p = doc.add_paragraph()
                _apply_rtl_p(p, space_before=12, space_after=4)
                _add_run_rtl(p, subtitle_text, size_pt=12.5, bold=True, color_rgb=color)

            def _add_callout(text, prefix=None, border_color="002045", fill_hex="F8FAFC", is_bold=False):
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tblPr = tbl._tbl.tblPr
                if tblPr.find(qn("w:bidiVisual")) is None:
                    tblPr.append(OxmlElement("w:bidiVisual"))
                cell = tbl.cell(0, 0)
                tcPr = cell._tc.get_or_add_tcPr()
                if fill_hex:
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), fill_hex)
                    tcPr.append(shd)

                tcMar = OxmlElement("w:tcMar")
                for edge, val in [("top", 120), ("bottom", 120), ("left", 180), ("right", 180)]:
                    m = OxmlElement(f"w:{edge}")
                    m.set(qn("w:w"), str(val))
                    m.set(qn("w:type"), "dxa")
                    tcMar.append(m)
                tcPr.append(tcMar)

                tcBorders = OxmlElement("w:tcBorders")
                r_bdr = OxmlElement("w:right")
                r_bdr.set(qn("w:val"), "single")
                r_bdr.set(qn("w:sz"), "24")  # 3 pt
                r_bdr.set(qn("w:color"), border_color)
                tcBorders.append(r_bdr)
                for edge in ["top", "bottom", "left"]:
                    b = OxmlElement(f"w:{edge}")
                    b.set(qn("w:val"), "none")
                    tcBorders.append(b)
                tcPr.append(tcBorders)

                p = cell.paragraphs[0]
                _apply_rtl_p(p, space_before=0, space_after=0)
                if prefix:
                    _add_run_rtl(p, prefix + " ", size_pt=11, bold=True, color_rgb=TEKHELET)
                _add_run_rtl(p, text, size_pt=11, bold=is_bold, italic=False, color_rgb=CHARCOAL)

                spacer = doc.add_paragraph()
                _apply_rtl_p(spacer, space_before=0, space_after=4)

            # Title
            title_p = doc.add_paragraph()
            _apply_rtl_p(title_p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
            _add_run_rtl(title_p, f"חוברת ליווי לדף מקורות: {self.title or self.topic}", size_pt=18, bold=True, color_rgb=TEKHELET)

            topic_p = doc.add_paragraph()
            _apply_rtl_p(topic_p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=14)
            _add_run_rtl(topic_p, f"נושא הסוגיה: {self.topic}", size_pt=12, bold=True, color_rgb=GOLD)
            _add_run_rtl(topic_p, "  ·  Chavruta.AI בית מדרש יוצר", size_pt=9.5, italic=True, color_rgb=MUTED)

            # Core inquiry
            _add_heading_1("שאלת היסוד וציר החקירה")
            _add_callout(self.core_inquiry, prefix="שאלת המפתח:", border_color="B8860B", fill_hex="FDFBF7", is_bold=True)

            # Sugya Flowchart / Outline
            if self.flowchart_mermaid:
                _add_heading_1("מפת מהלך הסוגיה (תרשים זרימה)")
                clean_lines = [
                    line.strip()
                    for line in self.flowchart_mermaid.split("\n")
                    if line.strip() and not line.strip().startswith("flowchart") and not line.strip().startswith("graph")
                ]
                if clean_lines:
                    tbl = doc.add_table(rows=0, cols=1)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tblPr = tbl._tbl.tblPr
                    if tblPr.find(qn("w:bidiVisual")) is None:
                        tblPr.append(OxmlElement("w:bidiVisual"))
                    for idx, line in enumerate(clean_lines):
                        step_text = line.replace("-->", "  ←  ").replace("---", "  —  ").replace("==>", "  ⇐  ")
                        for ch in ["[", "]", "(", ")", "{", "}"]:
                            step_text = step_text.replace(ch, "")
                        row = tbl.add_row()
                        c = row.cells[0]
                        tcPr = c._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:fill"), "F8FAFC" if idx % 2 == 0 else "FFFFFF")
                        tcPr.append(shd)
                        p = c.paragraphs[0]
                        _apply_rtl_p(p, space_before=3, space_after=3)
                        _add_run_rtl(p, f"• {step_text}", size_pt=10.5, color_rgb=NAVY_ACCENT)
                    spacer = doc.add_paragraph()
                    _apply_rtl_p(spacer, space_before=0, space_after=4)

            # Sections
            _add_heading_1("ביאור מקורות הדף")
            for s in self.sections:
                display_ref = (hebrew_display_ref(s.ref) or s.ref) if s.ref else ""
                status_text = STATUS_LABELS_HE.get(s.status, s.status)

                _add_heading_2(f"מקור {s.index}: {s.title} ({s.role_tag})")

                meta_p = doc.add_paragraph()
                _apply_rtl_p(meta_p, space_before=0, space_after=3)
                if display_ref:
                    _add_run_rtl(meta_p, f"מראה מקום: {display_ref}  |  ", size_pt=10, bold=True, color_rgb=NAVY_ACCENT)
                _add_run_rtl(
                    meta_p,
                    f"סטטוס: {status_text}",
                    size_pt=9.5,
                    italic=True,
                    color_rgb=MUTED if s.status != STATUS_MISSING_REF else DARK_RED,
                )

                if s.status == STATUS_MISSING_REF:
                    _add_callout(
                        "מראה מקום זה אינו מצוטט בדף ואינו קיים במאגר. לא ניתן לבאר ללא ציטוט הטקסט — באפשרותך להזין את לשונו בצ'אט.",
                        prefix="⚠️ שים לב:",
                        border_color="D97706",
                        fill_hex="FEF3C7",
                    )
                    continue

                if s.source_snippet:
                    _add_callout(s.source_snippet, prefix="לשון המקור:", border_color="002045", fill_hex="F8FAFC", is_bold=True)

                if s.plain_explanation:
                    exp_p = doc.add_paragraph()
                    _apply_rtl_p(exp_p, space_before=2, space_after=3)
                    _add_run_rtl(exp_p, "ביאור הפשט והמהלך: ", size_pt=11, bold=True, color_rgb=NAVY_ACCENT)
                    _add_run_rtl(exp_p, s.plain_explanation, size_pt=11, color_rgb=CHARCOAL)

                if s.diyuk:
                    diyuk_p = doc.add_paragraph()
                    _apply_rtl_p(diyuk_p, space_before=2, space_after=3)
                    _add_run_rtl(diyuk_p, "נקודת הדיוק: ", size_pt=11, bold=True, color_rgb=GOLD)
                    _add_run_rtl(diyuk_p, s.diyuk, size_pt=11, italic=True, color_rgb=CHARCOAL)

                if s.difficult_words:
                    words_p = doc.add_paragraph()
                    _apply_rtl_p(words_p, space_before=2, space_after=3)
                    _add_run_rtl(words_p, "ביאורי מילים: ", size_pt=10.5, bold=True, color_rgb=NAVY_ACCENT)
                    words_str = "  |  ".join(f"{k} — {v}" for k, v in s.difficult_words.items())
                    _add_run_rtl(words_p, words_str, size_pt=10, color_rgb=CHARCOAL)

                if s.author_note:
                    note_p = doc.add_paragraph()
                    _apply_rtl_p(note_p, space_before=2, space_after=6)
                    _add_run_rtl(note_p, "💡 הערת עורך הדף: ", size_pt=10.5, bold=True, color_rgb=GOLD)
                    _add_run_rtl(note_p, s.author_note, size_pt=10.5, italic=True, color_rgb=CHARCOAL)

            # Opinion table
            if self.opinion_table:
                _add_heading_1("טבלת השוואת שיטות ומחלוקות")
                table = doc.add_table(rows=1, cols=4)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                tblPr = table._tbl.tblPr
                if tblPr.find(qn("w:bidiVisual")) is None:
                    tblPr.append(OxmlElement("w:bidiVisual"))

                col_widths = [Inches(1.6), Inches(2.2), Inches(1.5), Inches(1.5)]
                headers = ["שיטה / בעל הדעה", "סברת היסוד", "הראיה / המקור", "נפקא מינה"]
                hdr_cells = table.rows[0].cells
                for idx, text in enumerate(headers):
                    hdr_cells[idx].width = col_widths[idx]
                    tcPr = hdr_cells[idx]._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), "1C355E")
                    tcPr.append(shd)

                    tcMar = OxmlElement("w:tcMar")
                    for edge, val in [("top", 120), ("bottom", 120), ("left", 140), ("right", 140)]:
                        m = OxmlElement(f"w:{edge}")
                        m.set(qn("w:w"), str(val))
                        m.set(qn("w:type"), "dxa")
                        tcMar.append(m)
                    tcPr.append(tcMar)

                    p = hdr_cells[idx].paragraphs[0]
                    _apply_rtl_p(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=2, space_after=2)
                    _add_run_rtl(p, text, size_pt=10.5, bold=True, color_rgb=RGBColor(255, 255, 255))

                for r_idx, row_data in enumerate(self.opinion_table):
                    row = table.add_row()
                    row_cells = row.cells
                    vals = [
                        row_data.get("opinion", ""),
                        row_data.get("reason", ""),
                        row_data.get("proof", ""),
                        row_data.get("nafka_mina", ""),
                    ]
                    fill_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                    for c_idx, val in enumerate(vals):
                        row_cells[c_idx].width = col_widths[c_idx]
                        tcPr = row_cells[c_idx]._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:fill"), fill_color)
                        tcPr.append(shd)

                        tcMar = OxmlElement("w:tcMar")
                        for edge, val_pad in [("top", 100), ("bottom", 100), ("left", 140), ("right", 140)]:
                            m = OxmlElement(f"w:{edge}")
                            m.set(qn("w:w"), str(val_pad))
                            m.set(qn("w:type"), "dxa")
                            tcMar.append(m)
                        tcPr.append(tcMar)

                        tcBorders = OxmlElement("w:tcBorders")
                        for edge in ["top", "left", "bottom", "right"]:
                            b = OxmlElement(f"w:{edge}")
                            b.set(qn("w:val"), "single")
                            b.set(qn("w:sz"), "4")  # 0.5 pt
                            b.set(qn("w:color"), "E2E8F0")
                            tcBorders.append(b)
                        tcPr.append(tcBorders)

                        p = row_cells[c_idx].paragraphs[0]
                        _apply_rtl_p(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=2, space_after=2)
                        _add_run_rtl(p, val, size_pt=10, bold=(c_idx == 0), color_rgb=CHARCOAL)

                spacer = doc.add_paragraph()
                _apply_rtl_p(spacer, space_before=0, space_after=6)

            # Chavruta questions
            if self.chavruta_questions:
                _add_heading_1("שאלות לעיון וחזרת החברותא")
                q_cats = [
                    ("peshat", "1. שאלות פשט והבנת המקורות:"),
                    ("comparison", "2. שאלות השוואה ודיוק שיטות:"),
                    ("sevara", "3. שאלות סברא והעמקה מושגית:"),
                ]
                for key, cat_title in q_cats:
                    qs = self.chavruta_questions.get(key, [])
                    if qs:
                        _add_heading_2(cat_title, color=TEKHELET)
                        for q in qs:
                            qp = doc.add_paragraph()
                            _apply_rtl_p(qp, space_before=1, space_after=3)
                            _add_run_rtl(qp, "•  ", size_pt=11, bold=True, color_rgb=GOLD)
                            _add_run_rtl(qp, q, size_pt=11, color_rgb=CHARCOAL)

            # Summary
            if self.summary:
                _add_heading_1("סיכום ומסקנות הסוגיה")
                _add_callout(self.summary, prefix="סיכום הלכה למעשה:", border_color="002045", fill_hex="F8FAFC")

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as exc:
            _log.warning("docx generation failed: %s", exc)
            return b""

    def to_html_printable(self) -> str:
        """Produce an elegant, standalone printable HTML companion guide for A4 printing and PDF export."""
        title_esc = html.escape(self.title or self.topic or "דף מקורות")
        topic_esc = html.escape(self.topic or "")
        inquiry_esc = html.escape(self.core_inquiry or "")

        from chavruta.corpus.refs import hebrew_display_ref

        sections_html = []
        for s in self.sections:
            disp_ref = (hebrew_display_ref(s.ref) or s.ref) if s.ref else ""
            s_title = html.escape(s.title or f"מקור {s.index}")
            s_role = html.escape(s.role_tag or "מקור")

            if s.status == STATUS_VERIFIED_CORPUS:
                badge_class = "status-verified"
                badge_text = "מאומת מהמאגר"
            elif s.status == STATUS_USER_PROVIDED:
                badge_class = "status-user"
                badge_text = "מדף המקורות"
            else:
                badge_class = "status-missing"
                badge_text = "חסר מקור במאגר"

            ref_html = (
                f'<div class="source-ref-line"><strong>מראה מקום:</strong> {html.escape(disp_ref)}</div>'
                if disp_ref
                else ""
            )

            if s.status == STATUS_MISSING_REF:
                content_html = """
            <div class="missing-alert">
              ⚠️ <strong>הערת מערכת:</strong> מראה מקום זה אינו מצוטט בדף ואינו קיים במאגר. לא ניתן לבאר ללא ציטוט הטקסט — באפשרותך להזין את לשונו בצ'אט.
            </div>
            """
            else:
                snippet_html = (
                    f'<div class="source-snippet-box">{html.escape(s.source_snippet)}</div>'
                    if s.source_snippet
                    else ""
                )
                exp_html = (
                    f'<div class="section-field"><span class="field-label">ביאור הפשט והמהלך:</span> {html.escape(s.plain_explanation)}</div>'
                    if s.plain_explanation
                    else ""
                )
                diyuk_html = (
                    f'<div class="diyuk-box"><span class="diyuk-label">💡 נקודת הדיוק:</span> <em>{html.escape(s.diyuk)}</em></div>'
                    if s.diyuk
                    else ""
                )

                words_html = ""
                if s.difficult_words:
                    words_str = " &nbsp;|&nbsp; ".join(
                        f"<strong>{html.escape(k)}</strong>: {html.escape(v)}" for k, v in s.difficult_words.items()
                    )
                    words_html = f'<div class="words-glossary"><span class="field-label">ביאורי מילים:</span> {words_str}</div>'

                note_html = (
                    f'<div class="author-note-box">✍️ <strong>הערת עורך הדף:</strong> {html.escape(s.author_note)}</div>'
                    if s.author_note
                    else ""
                )

                content_html = f"{snippet_html}{exp_html}{diyuk_html}{words_html}{note_html}"

            card = f"""
        <div class="source-card avoid-break">
          <div class="source-header">
            <div class="source-title-wrap">
              <span class="source-index-badge">מקור {s.index}</span>
              <span class="source-title">{s_title}</span>
              <span class="role-pill">{s_role}</span>
            </div>
            <span class="status-pill {badge_class}">{badge_text}</span>
          </div>
          {ref_html}
          {content_html}
        </div>
        """
            sections_html.append(card)

        sections_block = "\n".join(sections_html)

        # Opinion table
        table_block = ""
        if self.opinion_table:
            rows_html = []
            for r in self.opinion_table:
                op = html.escape(r.get("opinion", ""))
                re_ = html.escape(r.get("reason", ""))
                pr = html.escape(r.get("proof", ""))
                nm = html.escape(r.get("nafka_mina", ""))
                rows_html.append(f"""
            <tr>
              <td class="opinion-name">{op}</td>
              <td>{re_}</td>
              <td>{pr}</td>
              <td>{nm}</td>
            </tr>
            """)
            table_block = f"""
        <h2 class="sec-heading avoid-break">טבלת השוואת שיטות ומחלוקות</h2>
        <div class="table-container avoid-break">
          <table class="opinion-table">
            <thead>
              <tr>
                <th>שיטה / בעל הדעה</th>
                <th>סברת היסוד</th>
                <th>הראיה / המקור</th>
                <th>נפקא מינה</th>
              </tr>
            </thead>
            <tbody>
              {''.join(rows_html)}
            </tbody>
          </table>
        </div>
        """

        # Chavruta questions
        questions_block = ""
        if self.chavruta_questions:
            q_cards = []
            mapping = [
                ("peshat", "1. שאלות פשט והבנת המקורות"),
                ("comparison", "2. שאלות השוואה ודיוק שיטות"),
                ("sevara", "3. שאלות סברא והעמקה מושגית"),
            ]
            for key, cat_title in mapping:
                qs = self.chavruta_questions.get(key, [])
                if qs:
                    li_items = "".join(f"<li>{html.escape(q)}</li>" for q in qs)
                    q_cards.append(f"""
                <div class="q-card avoid-break">
                  <div class="q-card-title">{cat_title}</div>
                  <ul class="q-list">
                    {li_items}
                  </ul>
                </div>
                """)
            if q_cards:
                questions_block = f"""
            <h2 class="sec-heading avoid-break">שאלות לעיון וחזרת החברותא</h2>
            <div class="questions-grid avoid-break">
              {''.join(q_cards)}
            </div>
            """

        # Diagram block
        diagram_block = ""
        mermaid_script = ""
        if self.flowchart_mermaid:
            mermaid_raw = self.flowchart_mermaid.strip()
            diagram_block = f"""
        <h2 class="sec-heading avoid-break">מפת מהלך הסוגיה</h2>
        <div class="diagram-container avoid-break">
          <div class="mermaid">
{mermaid_raw}
          </div>
        </div>
        """
            mermaid_script = """
  <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
  <script>
    document.addEventListener("DOMContentLoaded", function() {
      if (window.mermaid) {
        mermaid.initialize({
          startOnLoad: true,
          theme: 'neutral',
          securityLevel: 'loose',
          fontFamily: 'Assistant, Frank Ruhl Libre, sans-serif'
        });
      }
    });
  </script>
        """

        # Summary
        summary_block = ""
        if self.summary:
            summary_block = f"""
        <h2 class="sec-heading avoid-break">סיכום ומסקנות הסוגיה</h2>
        <div class="summary-card avoid-break">
          <p class="summary-text">{html.escape(self.summary)}</p>
        </div>
        """

        return f"""<!DOCTYPE html>
<html dir="rtl" lang="he">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>חוברת ליווי לדף מקורות — {title_esc}</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Assistant:wght@400;500;600;700&family=Frank+Ruhl+Libre:wght@400;500;600;700;800&family=David+Libre:wght@400;500;700&display=swap" rel="stylesheet">
  <style>
    :root {{
      --primary: #002045;
      --primary-light: #1C355E;
      --accent-gold: #B8860B;
      --accent-gold-light: #FDFBF7;
      --ink: #1C1A17;
      --muted: #64748B;
      --border-subtle: #E2E8F0;
      --bg-page: #F1F5F9;
      --bg-card: #FFFFFF;
    }}

    @page {{
      size: A4 portrait;
      margin: 15mm 12mm 15mm 12mm;
      @bottom-center {{
        content: counter(page);
        font-family: 'Assistant', sans-serif;
        font-size: 9pt;
        color: #64748B;
      }}
    }}

    * {{
      box-sizing: border-box;
      -webkit-print-color-adjust: exact !important;
      print-color-adjust: exact !important;
    }}

    body {{
      font-family: 'Frank Ruhl Libre', 'David Libre', 'David', serif;
      direction: rtl;
      background-color: var(--bg-page);
      color: var(--ink);
      margin: 0;
      padding: 24px;
      line-height: 1.65;
      font-size: 15.5px;
    }}

    .sheet-wrapper {{
      max-width: 880px;
      margin: 0 auto;
      background: var(--bg-card);
      padding: 44px 52px;
      box-shadow: 0 4px 20px rgba(0, 32, 69, 0.08);
      border-radius: 12px;
      border: 1px solid var(--border-subtle);
    }}

    /* Print toolbar (screen only) */
    .print-toolbar {{
      position: sticky;
      top: 16px;
      z-index: 100;
      max-width: 880px;
      margin: 0 auto 20px auto;
      background: rgba(0, 32, 69, 0.94);
      backdrop-filter: blur(8px);
      color: #ffffff;
      padding: 10px 20px;
      border-radius: 9999px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      box-shadow: 0 8px 24px rgba(0, 0, 0, 0.2);
      font-family: 'Assistant', sans-serif;
    }}
    .print-toolbar-info {{
      display: flex;
      align-items: center;
      gap: 8px;
      font-size: 13.5px;
      font-weight: 500;
    }}
    .btn-print {{
      background: linear-gradient(135deg, #B8860B 0%, #D4AF37 100%);
      color: #002045;
      font-family: 'Assistant', sans-serif;
      font-weight: 700;
      font-size: 14px;
      padding: 8px 22px;
      border: none;
      border-radius: 9999px;
      cursor: pointer;
      display: flex;
      align-items: center;
      gap: 6px;
      transition: transform 0.15s ease, opacity 0.15s ease;
      box-shadow: 0 2px 8px rgba(184, 134, 11, 0.3);
    }}
    .btn-print:hover {{
      opacity: 0.95;
      transform: scale(1.02);
    }}

    /* Print media rules */
    @media print {{
      body {{
        background: #ffffff;
        padding: 0;
      }}
      .sheet-wrapper {{
        max-width: 100%;
        padding: 0;
        box-shadow: none;
        border: none;
        border-radius: 0;
      }}
      .no-print {{
        display: none !important;
      }}
      .avoid-break {{
        break-inside: avoid;
        page-break-inside: avoid;
      }}
      .page-break {{
        break-before: page;
        page-break-before: always;
      }}
    }}

    /* Header */
    .sheet-header {{
      text-align: center;
      border-bottom: 2px solid var(--accent-gold);
      padding-bottom: 18px;
      margin-bottom: 26px;
    }}
    .badge-top {{
      display: inline-block;
      font-family: 'Assistant', sans-serif;
      font-size: 12px;
      font-weight: 700;
      letter-spacing: 0.05em;
      color: var(--accent-gold);
      background: #FDFBF7;
      border: 1px solid rgba(184, 134, 11, 0.3);
      padding: 3px 14px;
      border-radius: 9999px;
      margin-bottom: 8px;
    }}
    .sheet-title {{
      font-size: 27px;
      font-weight: 800;
      color: var(--primary);
      margin: 0 0 6px 0;
      line-height: 1.3;
    }}
    .sheet-topic {{
      font-size: 16px;
      font-weight: 600;
      color: var(--primary-light);
      margin: 0;
    }}

    /* Headings */
    h2.sec-heading {{
      font-size: 19px;
      font-weight: 700;
      color: var(--primary);
      border-bottom: 1.5px solid var(--accent-gold);
      padding-bottom: 5px;
      margin: 28px 0 14px 0;
    }}

    /* Core Inquiry */
    .inquiry-card {{
      background: var(--accent-gold-light);
      border-right: 5px solid var(--accent-gold);
      border-radius: 8px 0 0 8px;
      padding: 16px 20px;
      margin-bottom: 22px;
      border-top: 1px solid #F1E9D7;
      border-bottom: 1px solid #F1E9D7;
      border-left: 1px solid #F1E9D7;
    }}
    .inquiry-label {{
      font-family: 'Assistant', sans-serif;
      font-size: 12.5px;
      font-weight: 700;
      color: var(--accent-gold);
      text-transform: uppercase;
      margin-bottom: 4px;
    }}
    .inquiry-text {{
      font-size: 16.5px;
      font-weight: 600;
      color: var(--primary);
      line-height: 1.55;
      margin: 0;
    }}

    /* Diagram */
    .diagram-container {{
      background: #FAFBFD;
      border: 1px solid var(--border-subtle);
      border-radius: 8px;
      padding: 18px;
      margin: 14px 0 22px 0;
      text-align: center;
      overflow-x: auto;
    }}

    /* Sources */
    .source-card {{
      background: #FFFFFF;
      border: 1px solid var(--border-subtle);
      border-radius: 8px;
      padding: 18px 20px;
      margin-bottom: 18px;
      box-shadow: 0 2px 6px rgba(0,0,0,0.02);
    }}
    .source-header {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      flex-wrap: wrap;
      gap: 10px;
      margin-bottom: 10px;
      padding-bottom: 8px;
      border-bottom: 1px solid #F1F5F9;
    }}
    .source-title-wrap {{
      display: flex;
      align-items: center;
      gap: 8px;
    }}
    .source-index-badge {{
      background: var(--primary);
      color: #ffffff;
      font-family: 'Assistant', sans-serif;
      font-size: 11.5px;
      font-weight: 700;
      padding: 2px 8px;
      border-radius: 5px;
    }}
    .source-title {{
      font-size: 16.5px;
      font-weight: 700;
      color: var(--primary);
    }}
    .role-pill {{
      font-family: 'Assistant', sans-serif;
      font-size: 11.5px;
      font-weight: 600;
      color: var(--primary-light);
      background: #EDF2F7;
      padding: 2px 8px;
      border-radius: 5px;
    }}
    .status-pill {{
      font-family: 'Assistant', sans-serif;
      font-size: 11px;
      font-weight: 600;
      padding: 2px 9px;
      border-radius: 9999px;
    }}
    .status-verified {{
      background: #ECFDF5;
      color: #065F46;
      border: 1px solid #A7F3D0;
    }}
    .status-user {{
      background: #F1F5F9;
      color: #334155;
      border: 1px solid #CBD5E1;
    }}
    .status-missing {{
      background: #FEF3C7;
      color: #92400E;
      border: 1px solid #FDE68A;
    }}

    .source-ref-line {{
      font-family: 'Assistant', sans-serif;
      font-size: 13px;
      color: var(--muted);
      margin-bottom: 10px;
    }}

    .source-snippet-box {{
      background: #F8FAFC;
      border-right: 4px solid var(--primary);
      padding: 10px 14px;
      border-radius: 4px 0 0 4px;
      font-size: 15.5px;
      font-weight: 600;
      color: var(--ink);
      margin-bottom: 12px;
      line-height: 1.6;
    }}

    .section-field {{
      margin-bottom: 8px;
      font-size: 15px;
    }}
    .field-label {{
      font-weight: 700;
      color: var(--primary-light);
      margin-left: 6px;
    }}
    .diyuk-box {{
      background: #FFFDF5;
      border-right: 3px solid var(--accent-gold);
      padding: 8px 12px;
      border-radius: 4px 0 0 4px;
      margin-bottom: 8px;
      font-size: 14.5px;
    }}
    .diyuk-label {{
      font-weight: 700;
      color: #B45309;
      margin-left: 6px;
    }}
    .words-glossary {{
      background: #F8FAFC;
      padding: 6px 12px;
      border-radius: 6px;
      font-size: 13.5px;
      margin-bottom: 8px;
    }}
    .author-note-box {{
      background: #F8FAFC;
      border: 1px dashed var(--accent-gold);
      padding: 6px 12px;
      border-radius: 6px;
      font-size: 13.5px;
      color: var(--primary-light);
      margin-top: 6px;
    }}
    .missing-alert {{
      background: #FFFBEB;
      border: 1px solid #FDE68A;
      color: #92400E;
      padding: 10px 14px;
      border-radius: 6px;
      font-size: 13.5px;
    }}

    /* Table */
    .table-container {{
      overflow-x: auto;
      margin-top: 10px;
      margin-bottom: 20px;
    }}
    table.opinion-table {{
      width: 100%;
      border-collapse: collapse;
      text-align: right;
      font-size: 14px;
    }}
    table.opinion-table th {{
      background: var(--primary-light);
      color: #FFFFFF;
      font-family: 'Assistant', sans-serif;
      font-weight: 700;
      padding: 9px 12px;
      border: 1px solid var(--primary-light);
    }}
    table.opinion-table td {{
      padding: 9px 12px;
      border: 1px solid var(--border-subtle);
      vertical-align: top;
    }}
    table.opinion-table tr:nth-child(even) {{
      background-color: #F8FAFC;
    }}
    table.opinion-table td.opinion-name {{
      font-weight: 700;
      color: var(--primary);
      width: 20%;
    }}

    /* Questions */
    .questions-grid {{
      display: flex;
      flex-direction: column;
      gap: 12px;
      margin-top: 12px;
      margin-bottom: 20px;
    }}
    .q-card {{
      background: #FAFBFD;
      border: 1px solid var(--border-subtle);
      border-radius: 8px;
      padding: 14px 18px;
    }}
    .q-card-title {{
      font-family: 'Assistant', sans-serif;
      font-size: 14.5px;
      font-weight: 700;
      color: var(--primary);
      margin-bottom: 8px;
    }}
    .q-list {{
      margin: 0;
      padding-right: 18px;
    }}
    .q-list li {{
      margin-bottom: 4px;
      color: var(--ink);
      font-size: 14.5px;
    }}

    /* Summary */
    .summary-card {{
      background: #FDFBF7;
      border: 1px solid #EADDC6;
      border-right: 5px solid var(--accent-gold);
      border-radius: 8px 0 0 8px;
      padding: 16px 20px;
      margin-top: 14px;
      margin-bottom: 20px;
    }}
    .summary-text {{
      margin: 0;
      font-size: 15px;
      line-height: 1.6;
    }}

    /* Footer */
    .sheet-footer {{
      text-align: center;
      border-top: 1px solid var(--border-subtle);
      padding-top: 14px;
      margin-top: 32px;
      font-family: 'Assistant', sans-serif;
      font-size: 12px;
      color: var(--muted);
    }}
  </style>
</head>
<body>
  <!-- Print Action Toolbar (Screen Only) -->
  <div class="print-toolbar no-print">
    <div class="print-toolbar-info">
      <span>📄 <strong>דף מקורות מוכן להדפסה</strong> — לחץ להדפסה ישירה או לשמירה כקובץ PDF איכותי</span>
    </div>
    <button onclick="window.print()" class="btn-print">
      🖨️ הדפס / שמור כ-PDF
    </button>
  </div>

  <div class="sheet-wrapper">
    <header class="sheet-header">
      <div class="badge-top">CHAVRUTA.AI · בית מדרש יוצר</div>
      <h1 class="sheet-title">חוברת ליווי לדף מקורות: {title_esc}</h1>
      <p class="sheet-topic">נושא הסוגיה: {topic_esc}</p>
    </header>

    <div class="inquiry-card avoid-break">
      <div class="inquiry-label">שאלת היסוד וציר החקירה</div>
      <p class="inquiry-text">{inquiry_esc}</p>
    </div>

    {diagram_block}

    <h2 class="sec-heading">ביאור מקורות הדף</h2>
    {sections_block}

    {table_block}

    {questions_block}

    {summary_block}

    <footer class="sheet-footer">
      <div>הופק ע״י Chavruta.AI — פלטפורמת לימוד תורני מבוססת מקורות</div>
      <div style="margin-top: 4px; font-size: 11px; opacity: 0.8;">עוצב להדפסה מושלמת בפורמט A4</div>
    </footer>
  </div>

  {mermaid_script}
</body>
</html>
"""


# ── XML-Isolated Context Construction (Principle I) ──────────────────────────

def build_sourcesheet_prompt_context(
    items: list[ParsedSourceItem],
    corpus_lookup: dict[str, str] | None = None,
) -> str:
    """Build strictly-isolated XML source blocks to prevent parametric hallucinations."""
    blocks = []
    lookup = corpus_lookup or {}

    for item in items:
        source_id = f"S{item.index}"
        ref = item.ref or "None"
        canonical = item.canonical_sefaria_ref or ""
        corpus_text = lookup.get(canonical) or lookup.get(ref)

        has_substantive_body = bool(
            item.cleaned_text
            and item.cleaned_text.strip() != item.header.strip()
            and len(item.cleaned_text.strip()) > 5
        )
        if corpus_text:
            status = STATUS_VERIFIED_CORPUS
        elif has_substantive_body:
            status = STATUS_USER_PROVIDED
        else:
            status = STATUS_MISSING_REF

        block = [f'<source id="{source_id}" status="{status}" ref="{ref}">']
        block.append(f"  <header>{item.header}</header>")
        if item.dibur_hamatchil:
            block.append(f"  <dibur_hamatchil>{item.dibur_hamatchil}</dibur_hamatchil>")
        if item.author_note_text:
            block.append(f"  <author_annotation>{item.author_note_text}</author_annotation>")
        if corpus_text:
            block.append(f"  <corpus_verified_text>{corpus_text}</corpus_verified_text>")
        elif status == STATUS_USER_PROVIDED:
            block.append(f"  <user_provided_text>{item.cleaned_text}</user_provided_text>")
        else:
            block.append("  <instruction>UNINDEXED BARE REF. DO NOT HALLUCINATE TEXT.</instruction>")
        block.append("</source>")
        blocks.append("\n".join(block))

    return "\n\n".join(blocks)


# ── Structured Sugya Synthesis Engine ────────────────────────────────────────

def _clean_topic_text(text: str) -> str:
    """Sanitize and strip leaked system instructions, filenames or imperative query prefixes from topic."""
    if not text:
        return ""
    cleaned = text
    # Strip system prompt leakage
    cleaned = re.sub(r"\(לא מהמאגר[^\)]*\)", "", cleaned)
    cleaned = re.sub(r"\(not from the corpus[^\)]*\)", "", cleaned)
    cleaned = re.sub(r"##\s+(?:מקורות שצירף המשתמש|Sources the user attached)[^\n]*", "", cleaned)
    cleaned = re.sub(r"###\s+(?:מקור|Source)\s*\d*", "", cleaned)
    # Strip markdown heading and file extension artifacts (e.g. "### 02 הרב חיים וולפסון - פרישת כהן גדול.docx")
    cleaned = re.sub(r"###\s*", "", cleaned)
    cleaned = re.sub(r"\.(?:docx|pdf|txt|doc)\b", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^\s*\d+[\s._-]*", "", cleaned)
    # Strip author names/prefixes like "הרב חיים וולפסון - "
    cleaned = re.sub(r"^(?:הרב\s+[^\n-–—]+[-–—]\s*)", "", cleaned)
    # Strip imperative query prefixes
    imperative_prefixes = [
        r"^(?:תסכם|סכם|באר|תבאר|הסבר|תסביר|נתח|תנתח|בנה|תבנה|הצג|תציג)\s+(?:לי\s+)?(?:את\s+)?(?:כל\s+)?(?:המהלך\s+של\s+)?(?:דף\s+המקורות|הסוגיה|המקורות|הדף)?(?:\s*[:—–-])?\s*",
        r"^(?:מה\s+הוא|מהו|מהם|כיצד|איך)\s+",
    ]
    for pat in imperative_prefixes:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE).strip()
    return cleaned.strip().strip(":—–- ")


def _clean_hebrew_prose(text: str) -> str:
    """Strip foreign script artifacts and replace transliterations with pure Hebrew."""
    if not text:
        return ""
    cleaned = re.sub(r'[\u4e00-\u9fff]+', '', text)
    cleaned = re.sub(r'\britual[יi]ת?\b', 'של טהרה', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\britual[יi]ם\b', 'דיני טהרה', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\britual\b', 'טהרה וקדושה', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\b[rR]itual\b', 'טהרה וקדושה', cleaned)
    return cleaned.strip()


def _synthesize_with_llm(
    items: list[ParsedSourceItem],
    topic_hint: str,
    corpus_lookup: dict[str, str],
    llm: Any,
    lang: str,
    user_instruction: str = "",
) -> CompanionGuide | None:
    """Invoke the LLM to perform deep Torah synthesis of the source sheet."""
    sources_xml = build_sourcesheet_prompt_context(items, corpus_lookup)
    prompt = f"""אתה תלמיד חכם מובהק, מגיד שיעור ועורך תורני מומחה במערכת 'חברותא AI'.
לפניך דף מקורות תורני שחולץ לתוך מקטעי XML (מקורות מאומתים מהמאגר או טקסטים מדף המשתמש).
עליך לנתח את הדף ברמה למדנית ופדגוגית גבוהה, ולבנות חוברת ליווי מקיפה למהלך הסוגיה.

הנחיות חמורות (Principle I):
- התבסס אך ורק על המקורות המופיעים ב-XML למטה. אל תמציא מקורות או מובאות שלא ניתנו.
- אם מופיע מקור ללא טקסט (UNINDEXED BARE REF), ציין שהוא מראה מקום בלבד ואל תבדה את תוכנו.
- זהה את נושא הסוגיה האמיתי (למשל: "פרישת כהן גדול ביום הכיפורים", "מצות תלמוד תורה וגדריה", "ייאוש שלא מדעת"). אל תשתמש בשמות קבצים, שמות מרצים או במחרוזות מערכת כנושא.
- שפה וסגנון: כתוב אך ורק בעברית צחה, תורנית ועשירה. אין להשתמש באותיות לועזיות (למשל אל תכתוב R"A אלא ר' אבהו), אין להשתמש במילים לועזיות (כגון 'ריטואל' -> השתמש ב'מעשה המצוה' או 'טהרה'), ואין לשלב תווים בשפות זרות.

החזר פלט מובנה בפורמט JSON בלבד (עטוף ב-```json ... ``` או JSON ישיר בלבד, ללא מלל נוסף לפני או אחרי):
{{
  "topic": "נושא הסוגיה המרכזי והמדויק (2-5 מילים)",
  "core_inquiry": "שאלת היסוד, ציר החקירה והסברא העומדת במוקד הסוגיה (2-4 משפטים)",
  "summary": "סיכום מקיף, תמציתי ובהיר של מהלך הסוגיה, השתלשלות השיטות והמסקנה (3-5 פסקאות)",
  "sections": [
    {{
      "index": 1,
      "title": "כותרת המקור בעברית (למשל: בבא מציעא דף כ\"א ע\"א / רמב\"ם הלכות ת\"ת)",
      "role_tag": "מקור יסוד / עובדא דש\"ס | קושיא / דיוק | חידוש / יסוד הסברא | ראיה / סייעתא | שיטה חולקת | הכרעה הלכתית",
      "plain_explanation": "ביאור תמציתי ומאיר עיניים של המקור בעברית תורנית, תוך הדגשת מקומו במהלך",
      "diyuk": "דיוק בלשון המקור או דיבור המתחיל (אם ישנו)",
      "difficult_words": {{"מילה קשה": "פירושה"}}
    }}
  ],
  "opinion_table": [
    {{
      "opinion": "שם השיטה / בעל הדעה בעברית",
      "reason": "טעם וסברא מרכזית",
      "proof": "מקור וראיה",
      "nafka_mina": "נפקא מינה להלכה או למעשה"
    }}
  ],
  "chavruta_questions": {{
    "peshat": ["שאלת פשט והבנה 1", "שאלת פשט והבנה 2"],
    "comparison": ["שאלת השוואת שיטות 1", "שאלת השוואת שיטות 2"],
    "sevara": ["שאלת סברא, חקירה ולמדנות 1", "שאלת סברא ולמדנות 2"]
  }},
  "flowchart_mermaid": "flowchart TD\\n    A[שאלת היסוד] --> B[מקור 1]\\n    B --> C[מקור 2]"
}}

מקורות הדף לעיון:
{sources_xml}
"""
    try:
        response_text = ""
        system = "אתה תלמיד חכם מובהק, מגיד שיעור ועורך תורני מומחה במערכת 'חברותא AI'. כתוב בעברית תורנית צחה בלבד."
        if hasattr(llm, "generate"):
            grounded_prompt = GroundedPrompt(system=system, sources=[], question=prompt, bare=True)
            res = llm.generate(grounded_prompt, lang="he", max_tokens=3500, temperature=0.2)
            response_text = res.text if hasattr(res, "text") else str(res)
        elif hasattr(llm, "complete"):
            response_text = str(llm.complete(prompt))

        if not response_text:
            return None

        # Extract JSON block
        json_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", response_text, re.DOTALL)
        raw_json = json_match.group(1) if json_match else response_text.strip()
        if not json_match and "{" in raw_json and "}" in raw_json:
            raw_json = raw_json[raw_json.find("{"):raw_json.rfind("}") + 1]

        # Convert unescaped Hebrew acronym quotes (e.g. רמב"ם -> רמב״ם) to avoid invalid JSON syntax
        raw_json = re.sub(r'(?<=[\u0590-\u05FF])"(?=[\u0590-\u05FF])', '״', raw_json)

        data = json.loads(raw_json)
        raw_topic = data.get("topic") or topic_hint or (items[0].header if items else "סוגיה תורנית")
        topic = _clean_topic_text(raw_topic) or "סוגיה תורנית"
        core_inquiry = _clean_hebrew_prose(data.get("core_inquiry") or f"בירור יסודות וגדרי {topic}.")
        summary = _clean_hebrew_prose(data.get("summary") or "")

        llm_sections_data = {s.get("index", idx + 1): s for idx, s in enumerate(data.get("sections", []))}
        sections: list[SourceSection] = []
        citations: list[str] = []

        for item in items:
            sec_data = llm_sections_data.get(item.index, {})
            ref = item.ref or item.header
            canonical = item.canonical_sefaria_ref or ""
            corpus_text = corpus_lookup.get(canonical) or corpus_lookup.get(ref)
            has_body = bool(
                item.cleaned_text
                and item.cleaned_text.strip() != item.header.strip()
                and len(item.cleaned_text.strip()) > 5
            )

            if corpus_text:
                status = STATUS_VERIFIED_CORPUS
                snippet = corpus_text[:300] + ("…" if len(corpus_text) > 300 else "")
                citations.append(ref)
            elif has_body:
                status = STATUS_USER_PROVIDED
                snippet = item.cleaned_text
            else:
                status = STATUS_MISSING_REF
                snippet = "(טקסט אינו קיים בדף ואינו במאגר)"

            raw_title = sec_data.get("title") or item.header or ref or f"מקור {item.index}"
            title = _clean_topic_text(raw_title) or (ref or f"מקור {item.index}")
            role = _clean_hebrew_prose(sec_data.get("role_tag") or "מקור")
            explanation = _clean_hebrew_prose(sec_data.get("plain_explanation") or f"ביאור מקור {item.index} במסגרת מהלך הסוגיה.")
            diyuk = _clean_hebrew_prose(sec_data.get("diyuk") or (item.dibur_hamatchil and f'ד"ה "{item.dibur_hamatchil}"') or "")

            sec = SourceSection(
                index=item.index,
                title=title,
                ref=ref if ref != "None" else None,
                status=status,
                role_tag=role,
                source_snippet=snippet,
                expanded_context=corpus_text,
                plain_explanation=explanation,
                diyuk=diyuk or None,
                difficult_words=sec_data.get("difficult_words") or {},
                author_note=item.author_note_text,
            )
            sections.append(sec)

        flowchart = _clean_hebrew_prose(data.get("flowchart_mermaid") or "")
        opinion_table = [
            {k: _clean_hebrew_prose(v) for k, v in row.items()}
            for row in (data.get("opinion_table") or [])
        ]
        chavruta_questions = {
            level: [_clean_hebrew_prose(q) for q in qs]
            for level, qs in (data.get("chavruta_questions") or {}).items()
        }

        return CompanionGuide(
            title=f"מהלך הסוגיה — {topic}",
            topic=topic,
            core_inquiry=core_inquiry,
            flowchart_mermaid=flowchart,
            sections=sections,
            opinion_table=opinion_table,
            chavruta_questions=chavruta_questions,
            summary=summary,
            citations=citations,
        )
    except Exception as exc:
        _log.warning("sourcesheet LLM synthesis failed, falling back to deterministic: %s", exc)
        return None


# ── Structured Sugya Synthesis Engine ────────────────────────────────────────

def analyze_source_sheet(
    items: list[ParsedSourceItem],
    topic_hint: str = "",
    corpus_lookup: dict[str, str] | None = None,
    llm=None,
    lang: str = "he",
    user_instruction: str = "",
) -> CompanionGuide:
    """Synthesize a complete Source Sheet Companion Guide.

    Uses LLM for deep Torah analysis when available, with deterministic fallback.
    """
    if not items:
        return CompanionGuide(
            title="דף מקורות ריק",
            topic="ללא מקורות",
            core_inquiry="לא זוהו מקורות בדף שהועלה.",
            flowchart_mermaid="",
            sections=[],
            opinion_table=[],
            chavruta_questions={},
            summary="לא זוהו מקורות.",
        )

    lookup = corpus_lookup or {}

    # Try LLM synthesis first if LLM is provided
    if llm is not None:
        guide = _synthesize_with_llm(
            items=items,
            topic_hint=topic_hint,
            corpus_lookup=lookup,
            llm=llm,
            lang=lang,
            user_instruction=user_instruction,
        )
        if guide is not None:
            return guide

    # Deterministic fallback
    sections: list[SourceSection] = []
    citations: list[str] = []

    role_cycle = [
        "מקור יסוד / עובדא דש\"ס",
        "קושיא / דיוק",
        "חידוש / יסוד הסברא",
        "ראיה / סייעתא",
        "שיטה חולקת",
        "הכרעה הלכתית",
    ]

    for idx, item in enumerate(items):
        ref = item.ref or item.header
        canonical = item.canonical_sefaria_ref or ""
        corpus_text = lookup.get(canonical) or lookup.get(ref)

        has_substantive_body = bool(
            item.cleaned_text
            and item.cleaned_text.strip() != item.header.strip()
            and len(item.cleaned_text.strip()) > 5
        )
        if corpus_text:
            status = STATUS_VERIFIED_CORPUS
            snippet = corpus_text[:300] + ("…" if len(corpus_text) > 300 else "")
            citations.append(ref)
        elif has_substantive_body:
            status = STATUS_USER_PROVIDED
            snippet = item.cleaned_text
        else:
            status = STATUS_MISSING_REF
            snippet = "(טקסט אינו קיים בדף ואינו במאגר)"

        role = role_cycle[idx % len(role_cycle)]
        raw_title = item.header if item.header and len(item.header) < 60 else (ref or f"מקור {item.index}")
        title = _clean_topic_text(raw_title) or (ref or f"מקור {item.index}")

        sec = SourceSection(
            index=item.index,
            title=title,
            ref=ref if ref != "None" else None,
            status=status,
            role_tag=role,
            source_snippet=snippet,
            expanded_context=corpus_text,
            plain_explanation=f"ביאור מקור {item.index} במסגרת מהלך הסוגיה.",
            diyuk=item.dibur_hamatchil and f'עומד על הדיבור המתחיל "{item.dibur_hamatchil}"',
            author_note=item.author_note_text,
        )
        sections.append(sec)

    cleaned_topic = _clean_topic_text(topic_hint)
    first_header = _clean_topic_text(items[0].header) if items else ""
    detected_topic = cleaned_topic or first_header or (items[0].ref if items else "סוגיה תורנית")
    first_ref = items[0].ref or "סוגיית היסוד"

    # Mermaid Flowchart
    mermaid_nodes = []
    mermaid_nodes.append('flowchart TD')
    mermaid_nodes.append(f'    A["שאלת היסוד: {detected_topic[:40]}"] --> B["מקור 1: {first_ref}"]')
    for i in range(1, min(len(sections), 5)):
        prev_char = chr(ord('B') + i - 1)
        curr_char = chr(ord('B') + i)
        curr_sec = sections[i]
        mermaid_nodes.append(f'    {prev_char} -->|{curr_sec.role_tag}| {curr_char}["מקור {curr_sec.index}: {curr_sec.title[:30]}"]')
    flowchart_code = "\n".join(mermaid_nodes)

    opinion_table = [
        {
            "opinion": sections[0].title if sections else "שיטה ראשונה",
            "reason": "מקור הדין הבסיסי והעמדת המקרה",
            "proof": sections[0].ref or "ש\"ס",
            "nafka_mina": "הגדרת הדין היסודית",
        }
    ]
    if len(sections) > 1:
        opinion_table.append({
            "opinion": sections[1].title,
            "reason": "דיוק או קושיא על הפשט",
            "proof": sections[1].ref or "מפרשים",
            "nafka_mina": "צמצום או הרחבת גדרי הדין",
        })

    chavruta_questions = {
        "peshat": [
            f"מהי נקודת הפתיחה המרכזית העולה מתוך {sections[0].title}?",
            "הסבר את המושג המרכזי המוזכר במקור בלשונך.",
        ],
        "comparison": [
            f"מה ההבדל העיקרי בין דברי {sections[0].title} למקורות הבאים אחריו?",
            "איזו קושיא יישב המקור השני על המקור הראשון?",
        ],
        "sevara": [
            "האם ניתן להציע חילוק מושגי (חקירה) שיסביר את שתי הדרכים בסוגיה?",
            "כיצד משליכה הכרעה זו על מקרים מודרניים דומים?",
        ],
    }

    if user_instruction:
        summary_text = f"ניתוח דף המקורות בהתאם לבקשתך: {user_instruction}\n\nסוגיית {detected_topic} נפרסת לאורך {len(sections)} מקורות מרכזיים, המתווים את המעבר מן הדין היסודי אל הבירור המושגי וההכרעה."
    else:
        summary_text = f"סוגיית {detected_topic} נפרסת לאורך {len(sections)} מקורות מרכזיים, המתווים את המעבר מן הדין היסודי אל הבירור המושגי וההכרעה."

    return CompanionGuide(
        title=f"מהלך הסוגיה — {detected_topic}",
        topic=detected_topic,
        core_inquiry=f"הסוגיה עוסקת בבירור יסודות {detected_topic}, החל ממקורות השורש, דרך שיטות הראשונים והחקירות המרכזיות, ועד לבירור המסקנה וההשלכות המעשיות.",
        flowchart_mermaid=flowchart_code,
        sections=sections,
        opinion_table=opinion_table,
        chavruta_questions=chavruta_questions,
        summary=summary_text,
        citations=citations,
    )
