"""Unit tests for Source Sheet Analyzer (Spec 008 Phase 3)."""

from __future__ import annotations

import pytest

from chavruta.sourcesheet.analyzer import (
    STATUS_MISSING_REF,
    STATUS_USER_PROVIDED,
    STATUS_VERIFIED_CORPUS,
    analyze_source_sheet,
    build_sourcesheet_prompt_context,
)
from chavruta.sourcesheet.parser import parse_source_sheet


def test_analyze_source_sheet_and_markdown_export():
    raw = """
    1. בבא מציעא דף כ"א ע"א:
    אמר רבא: ייאוש שלא מדעת לא הוי ייאוש.

    2. רש"י שם ד"ה ייאוש שלא מדעת:
    כגון שנפלה ממנו אבידה ועדיין לא ידע שנפלה ממנו.

    3. שו"ת חתם סופר אה"ע סימן פ"ג
    """
    items = parse_source_sheet(raw)
    assert len(items) == 3

    corpus_lookup = {
        "Bava Metzia 21a": "משנה: אלו מציאות שלו ואלו חייב להכריז...",
    }

    guide = analyze_source_sheet(items, topic_hint="ייאוש שלא מדעת", corpus_lookup=corpus_lookup)

    assert guide.topic == "ייאוש שלא מדעת"
    assert len(guide.sections) == 3

    # Section 1: Verified in corpus
    assert guide.sections[0].status == STATUS_VERIFIED_CORPUS
    assert guide.sections[0].ref == "Bava Metzia 21a"

    # Section 2: User provided text (Rashi)
    assert guide.sections[1].status == STATUS_USER_PROVIDED

    # Section 3: Missing ref (no text and not in corpus)
    assert guide.sections[2].status == STATUS_MISSING_REF

    # Markdown export
    md = guide.to_markdown()
    assert "# חוברת ליווי לדף מקורות" in md
    assert "```mermaid" in md
    assert "טבלת השוואת שיטות" in md
    assert "שאלות לעיון וחזרת החברותא" in md


def test_build_sourcesheet_prompt_context_xml():
    raw = "1. בבא מציעא דף כ\"א ע\"א:\nאמר רבא ייאוש שלא מדעת."
    items = parse_source_sheet(raw)
    xml = build_sourcesheet_prompt_context(items, corpus_lookup={"Bava Metzia 21a": "טקסט מאומת"})
    assert '<source id="S1" status="VERIFIED_CORPUS"' in xml
    assert "<corpus_verified_text>טקסט מאומת</corpus_verified_text>" in xml


def test_sourcesheet_llm_synthesis_and_clean_topic():
    raw = "1. בבא מציעא דף כ\"א ע\"א:\nאמר רבא ייאוש שלא מדעת."
    items = parse_source_sheet(raw)

    class FakeLLMResponse:
        text = """```json
{
  "topic": "סוגיית ייאוש שלא מדעת",
  "core_inquiry": "האם ייאוש בעלים למפרע מועיל או שמא בעינן ידיעה בפועל.",
  "summary": "הסוגיה בבבא מציעא פותחת במחלוקת אביי ורבא...",
  "sections": [
    {
      "index": 1,
      "title": "בבא מציעא כ\"א ע\"א",
      "role_tag": "מקור יסוד / עובדא דש\"ס",
      "plain_explanation": "העמדת מחלוקת אביי ורבא בדין אבידה שנמצאה קודם שידעו הבעלים.",
      "diyuk": "מדייק רבא דכל שלא ידע לא הוי ייאוש",
      "difficult_words": {}
    }
  ],
  "opinion_table": [
    {"opinion": "רבא", "reason": "לא ידע לא מייאש", "proof": "משנה", "nafka_mina": "חייב להכריז"}
  ],
  "chavruta_questions": {
    "peshat": ["מהי סברת רבא?"],
    "comparison": [],
    "sevara": ["האם ייאוש הוא מעשה קניין או הסרת בעלות?"]
  },
  "flowchart_mermaid": "flowchart TD\\n    A --> B"
}
```"""

    class FakeLLM:
        def generate(self, prompt, **kwargs):
            return FakeLLMResponse()

    # Pass leaked system instruction as topic hint
    leaked_hint = "תסכם את הדף (לא מהמאגר — התייחס אליהם כמקור נוסף)"
    guide = analyze_source_sheet(items, topic_hint=leaked_hint, llm=FakeLLM())

    assert guide.topic == "סוגיית ייאוש שלא מדעת"
    assert "לא מהמאגר" not in guide.topic
    assert "סוגיית ייאוש שלא מדעת" in guide.title
    assert guide.core_inquiry == "האם ייאוש בעלים למפרע מועיל או שמא בעינן ידיעה בפועל."
    assert "העמדת מחלוקת אביי ורבא" in guide.sections[0].plain_explanation
    assert len(guide.opinion_table) == 1
    assert "מהי סברת רבא?" in guide.chavruta_questions["peshat"]


def test_sourcesheet_to_docx_bytes_styled_and_rtl():
    raw = """
    1. בבא מציעא דף כ"א ע"א:
    אמר רבא: ייאוש שלא מדעת לא הוי ייאוש.

    2. רש"י שם ד"ה ייאוש שלא מדעת:
    כגון שנפלה ממנו אבידה ועדיין לא ידע שנפלה ממנו.
    """
    items = parse_source_sheet(raw)
    guide = analyze_source_sheet(items, topic_hint="ייאוש שלא מדעת")

    docx_bytes = guide.to_docx_bytes()
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 5000
    assert docx_bytes.startswith(b"PK\x03\x04")

    import io
    import docx
    from docx.oxml.ns import qn

    doc = docx.Document(io.BytesIO(docx_bytes))
    assert len(doc.paragraphs) > 5
    assert len(doc.tables) >= 1

    # Check for RTL table setting (bidiVisual)
    found_bidi_table = any(t._tbl.tblPr.find(qn("w:bidiVisual")) is not None for t in doc.tables)
    assert found_bidi_table is True

    # Check for RTL paragraph setting (bidi)
    found_bidi_p = any(p._p.pPr is not None and p._p.pPr.find(qn("w:bidi")) is not None for p in doc.paragraphs)
    assert found_bidi_p is True

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "חוברת ליווי לדף מקורות" in full_text
    assert "ייאוש שלא מדעת" in full_text


def test_sourcesheet_to_html_printable():
    raw = """
    1. בבא מציעא דף כ"א ע"א:
    אמר רבא: ייאוש שלא מדעת לא הוי ייאוש.
    """
    items = parse_source_sheet(raw)
    guide = analyze_source_sheet(items, topic_hint="ייאוש שלא מדעת")

    html_content = guide.to_html_printable()
    assert "<!DOCTYPE html>" in html_content
    assert '<html dir="rtl" lang="he">' in html_content
    assert "@media print" in html_content
    assert "size: A4 portrait" in html_content
    assert "window.print()" in html_content
    assert "Frank Ruhl Libre" in html_content
    assert "ייאוש שלא מדעת" in html_content
    assert "שאלת היסוד וציר החקירה" in html_content
    assert "ביאור מקורות הדף" in html_content


